"""智研星枢 - 科研文件解析服务（多类型通用解析器）
支持: xlsx/xls/csv/pdf/docx/txt/md/py/json/xml/html/ipynb 等
"""
from app.utils.safe_json import safe_json_parse
import os
import csv
import json
import logging
from typing import Dict, Any, List
from app.agents.qwen_client import call_qwen

logger = logging.getLogger(__name__)

PARSE_PROMPT = """你是科研文件解析助手。请根据文件内容生成结构化摘要，包含：
1. 核心内容概述（≤200字）
2. 关键数据/变量/方法提取
3. 与学术研究的相关性评分（1-10）
4. 建议的后续分析方向
用JSON格式返回：{"summary": "...", "key_elements": [...], "relevance_score": N, "suggestions": [...]}"""


def _parse_csv(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        rows = [row for i, row in enumerate(reader) if i < 50]
    return "\n".join([",".join(row) for row in rows])


def _parse_excel(file_path: str) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        lines: List[str] = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"=== Sheet: {sheet_name} ===")
            row_count = 0
            for row in ws.iter_rows(values_only=True):
                if row_count >= 50:
                    lines.append("... (more rows omitted)")
                    break
                cells = [str(c) if c is not None else "" for c in row]
                lines.append("\t".join(cells))
                row_count += 1
        wb.close()
        return "\n".join(lines)[:8000]
    except ImportError:
        return "[Excel file, install openpyxl: pip install openpyxl]"
    except Exception as e:
        return f"[Excel parse error: {str(e)}]"


def _parse_pdf(file_path: str) -> str:
    try:
        import fitz
        doc = fitz.open(file_path)
        pages_text = []
        for page_num in range(min(len(doc), 10)):
            text = doc.load_page(page_num).get_text("text").strip()
            if text:
                pages_text.append(f"--- Page {page_num+1} ---\n{text}")
        doc.close()
        return "\n\n".join(pages_text)[:8000]
    except ImportError:
        return "[PDF file, install pymupdf: pip install pymupdf]"
    except Exception as e:
        return f"[PDF parse error: {str(e)}]"


def _parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                paragraphs.append("\t".join(cells))
        return "\n".join(paragraphs)[:8000]
    except ImportError:
        return "[Word file, install python-docx: pip install python-docx]"
    except Exception as e:
        return f"[Word parse error: {str(e)}]"


def _parse_ipynb(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            nb = json.load(f)
        lines: List[str] = []
        for cell in nb.get("cells", []):
            cell_type = cell.get("cell_type", "")
            source = "".join(cell.get("source", []))
            if source.strip():
                lines.append(f"[{cell_type}]\n{source}")
        return "\n\n".join(lines)[:8000]
    except Exception as e:
        return f"[Notebook parse error: {str(e)}]"


def _parse_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read(8000)


def _parse_image(file_path: str) -> str:
    size = os.path.getsize(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    return f"[Image {ext}, size: {size/1024:.1f}KB, use multimodal chat for image understanding]"


PARSER_MAP = {
    ".xlsx": _parse_excel, ".xls": _parse_excel,
    ".csv": _parse_csv, ".tsv": _parse_csv,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx, ".doc": _parse_docx,
    ".ipynb": _parse_ipynb,
    ".txt": _parse_text, ".md": _parse_text, ".py": _parse_text,
    ".json": _parse_text, ".xml": _parse_text, ".html": _parse_text,
    ".htm": _parse_text, ".yaml": _parse_text, ".yml": _parse_text,
    ".toml": _parse_text, ".ini": _parse_text, ".cfg": _parse_text,
    ".log": _parse_text, ".sql": _parse_text, ".sh": _parse_text,
    ".bat": _parse_text, ".r": _parse_text, ".R": _parse_text,
    ".java": _parse_text, ".c": _parse_text, ".cpp": _parse_text,
    ".h": _parse_text, ".js": _parse_text, ".ts": _parse_text,
    ".css": _parse_text, ".tex": _parse_text, ".bib": _parse_text,
    ".rst": _parse_text,
    ".png": _parse_image, ".jpg": _parse_image, ".jpeg": _parse_image,
    ".gif": _parse_image, ".bmp": _parse_image, ".webp": _parse_image,
    ".svg": _parse_text, ".tiff": _parse_image, ".tif": _parse_image,
}

SUPPORTED_EXTS = set(PARSER_MAP.keys())


async def parse_research_file(file_path: str, mime_type: str, description: str = "") -> Dict[str, Any]:
    ext = os.path.splitext(file_path)[1].lower()
    parser_fn = PARSER_MAP.get(ext)

    if parser_fn:
        try:
            content_preview = parser_fn(file_path)
        except Exception as e:
            logger.error(f"File parse failed [{ext}]: {e}")
            content_preview = f"[Parse failed: {str(e)}]"
    else:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(100)
            content_preview = _parse_text(file_path)
            logger.info(f"Unknown type {ext}, parsing as text")
        except (UnicodeDecodeError, Exception):
            size = os.path.getsize(file_path)
            content_preview = f"[Binary file {ext}, size: {size/1024:.1f}KB, not supported]"

    user_msg = f"File type: {ext}\nDescription: {description}\n\nContent preview:\n{content_preview[:6000]}"
    result = await call_qwen(PARSE_PROMPT, user_msg, model="qwen-max")

    parsed_data = {}
    try:
        start = result["content"].find("{")
        end = result["content"].rfind("}") + 1
        if start >= 0 and end > start:
            parsed_data = safe_json_parse(result["content"][start:end],fallback={},label="file_parser")
    except (json.JSONDecodeError, KeyError):
        parsed_data = {"summary": result.get("content", "")[:500]}

    return {
        "status": "success",
        "summary": parsed_data.get("summary", ""),
        "data": parsed_data,
        "tokens": result.get("tokens", {}).get("input", 0) + result.get("tokens", {}).get("output", 0),
    }